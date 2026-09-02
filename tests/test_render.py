#!/usr/bin/env python3
"""retainer-agreement 渲染回归测试（v1.4.0 补 v1.0.x「6 个下划线补丁」的欠账）。

双模式：pytest 可收集（test_* 函数、裸 assert），也可 `python3 tests/test_render.py` 直跑
（check.sh 门禁走直跑，不依赖 pytest——与 case-filing / civil-complaint-research 测试同款）。

覆盖：
 1. num2cn 文档示例（含 100000 → 壹拾万，docstring 曾写错为「拾万」）
 2. 红线词扫描正例（FEE_REDLINE_PATTERNS 六类逐个命中）
 3. 红线词扫描反例（canonical example 两段 + 样板句式 A/B 零命中）
 4. example 渲染版式（当事人四行：单 run [tab][值][tab] + 连续单下划线 + center 布局 sz28；
    律师费段金额下划线分段；contract_num 留空 → 5 空格占位）——v1.0.6 血泪史的回归锁
 5. 长值（30 CJK 住址）→ left 布局 + 自动缩字号（18 ≤ sz < 28）
 6. 含风险代理措辞渲染 → stderr 出 WARN [fee-redline] 且**不拦截**（exit 0、产物完好）
 7. canonical example 渲染 → stderr 无 fee-redline WARN
 v2.0.0 新增（manifest 驱动引擎 + 模板初始化）：
 8. golden 逐字节对照：固定日期 example → document.xml 与 tests/golden/zhuojian-minshang.document.xml 完全一致
    （v1.4.2 引擎产物基线；引擎重构不得改变卓建模板一个字节）
 9. 往返：渲染出的「已填合同」当样本 → init_template.py 建包 → 再渲染 → 仍与 golden 逐字节一致
    （证明初始化流程在真实卓建版式上可动态提炼变量 / 下划线 / tab 布局，不靠写死规则）
10. 合成通用模板（非卓建版式：tab leader 画线且值无下划线 / 行内下划线无 tab / 金额下划线收费段）：
    隐私清扫拦截残留样本值 → 补 placement 后通过 → 渲染：leader 行规整为单 run 下划线、行内下划线继承、
    金额下划线自动识别、样本值零残留
11. 零模板包 → exit 3 并提示初始化；缺必填字段 → exit 4 并列出字段名
12. 模板指纹漂移 → stderr WARN [template-drift] 不拦截；relock 后 WARN 消失
"""
import json
import subprocess
import sys
import tempfile
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "scripts"
EXAMPLE = Path(__file__).resolve().parents[1] / "examples" / "example.json"
sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

import num2cn  # noqa: E402
import render as R  # noqa: E402

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"
GOLDEN_XML = Path(__file__).resolve().parents[1] / "tests" / "golden" / "zhuojian-minshang.document.xml"
INIT = SCRIPTS_DIR / "init_template.py"
INSPECT = SCRIPTS_DIR / "inspect_template.py"
FIXED_DATE = {"cover_date_cn": "二〇二六年九月二日", "year": 2026}


def _fixed_example() -> dict:
    ex = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    ex.update(FIXED_DATE)
    return ex


def _document_xml(docx_path: Path) -> bytes:
    import zipfile
    with zipfile.ZipFile(docx_path) as z:
        return z.read("word/document.xml")


def _render_with(contract: dict, name: str, templates_dir: Path, extra_args=()):
    tmp = Path(tempfile.mkdtemp(prefix="retainer_test_"))
    cj = tmp / "contract.json"
    cj.write_text(json.dumps(contract, ensure_ascii=False), encoding="utf-8")
    out = tmp / name
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "render.py"), str(cj), str(out),
         "--templates-dir", str(templates_dir), *extra_args],
        capture_output=True, text=True,
    )
    return out, proc


def _init(docx: Path, mp: dict, out_dir: Path, *flags):
    tmp = Path(tempfile.mkdtemp(prefix="retainer_map_"))
    mpath = tmp / "map.json"
    mpath.write_text(json.dumps(mp, ensure_ascii=False), encoding="utf-8")
    return subprocess.run(
        [sys.executable, str(INIT), "init", "--docx", str(docx), "--map", str(mpath), "--out", str(out_dir), *flags],
        capture_output=True, text=True,
    )


def _render(contract: dict, name: str) -> tuple[Path, subprocess.CompletedProcess]:
    """写临时 contract.json → subprocess 跑 render.py（生产 CLI 形态）→ (docx路径, 进程)。"""
    tmp = Path(tempfile.mkdtemp(prefix="retainer_test_"))
    cj = tmp / "contract.json"
    cj.write_text(json.dumps(contract, ensure_ascii=False), encoding="utf-8")
    out = tmp / name
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "render.py"), str(cj), str(out)],
        capture_output=True, text=True,
    )
    return out, proc


def _client_row(doc, prefix):
    """按 label 前缀取当事人信息行，返回 (tabs定义, 值run列表)。"""
    for p in doc.paragraphs:
        if p.text.lstrip().startswith(prefix):
            pPr = p._p.find(qn('w:pPr'))
            tabs = pPr.find(qn('w:tabs'))
            tabdef = [(e.get(qn('w:val')), int(e.get(qn('w:pos')))) for e in tabs]
            vruns = [r for r in p._p.findall(qn('w:r')) if r.findall(qn('w:tab'))]
            return tabdef, vruns
    raise AssertionError(f"client row not found: {prefix}")


def _run_props(r_el):
    txt = ''.join(t.text or '' for t in r_el.findall(qn('w:t')))
    rPr = r_el.find(qn('w:rPr'))
    u = rPr.find(qn('w:u')) if rPr is not None else None
    sz = rPr.find(qn('w:sz')) if rPr is not None else None
    return txt, len(r_el.findall(qn('w:tab'))), \
        u.get(qn('w:val')) if u is not None else None, \
        int(sz.get(qn('w:val'))) if sz is not None else None


def test_num2cn_documented_examples():
    # docstring 例句逐个锁定；100000 曾在 docstring 误写「拾万」（代码一直对，本 task 顺修文档）
    cases = [(30000, '叁万'), (15000, '壹万伍仟'), (25800, '贰万伍仟捌佰'),
             (100000, '壹拾万'), (1000000, '壹佰万'), (108, '壹佰零捌'),
             (10080, '壹万零捌拾'), (0, '零')]
    for n, expect in cases:
        assert num2cn.int_to_cn(n) == expect, (n, num2cn.int_to_cn(n))
    assert num2cn.amount_to_cn(3000.50) == '叁仟元伍角整'


def test_redline_scan_positive():
    clauses = [
        '后期风险代理律师服务费按实现金额的15%计算。',            # 风险代理 + 提成类
        '若判决不利，律师费减免一半。',                            # 减免
        '判不离不收费。',                                          # 判不离
        '对方反悔的，本笔费用不再收取。',                          # 对方反悔
        '按回款比例支付律师费。',                                  # 按…比例
        '取得离婚证后才收取全部律师费。',                          # 结果挂钩
    ]
    hits = R.scan_fee_redlines(clauses)
    hit_labels = {(i, label) for i, label, _ in hits}
    assert (0, '风险代理') in hit_labels
    assert (1, '减免承诺') in hit_labels
    assert (2, '判不离不收') in hit_labels
    assert (3, '对方反悔不收') in hit_labels
    assert (4, '按比例提成') in hit_labels
    assert (5, '结果挂钩收费') in hit_labels


def test_redline_scan_negative():
    example = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    clean = list(example['fee_clauses']) + [
        '一审阶段：在签订本合同之日，甲方向乙方支付律师代理费人民币叁万元（¥30000）。',
        '执行阶段：在立案当日一次性支付律师服务费贰万元（¥20000）。',
        '自本合同签订之日起15日内未取得案件受理通知书的，本笔律师费甲方不再支付。',
    ]
    assert R.scan_fee_redlines(clean) == []


def test_render_example_layout():
    example = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    out, proc = _render(example, "example.docx")
    assert proc.returncode == 0, proc.stderr
    doc = Document(str(out))

    # 当事人四行：值 run 必须是**一个** run 内 [tab][值][tab]、整体单下划线（v1.0.6 根治双线/断线）
    for prefix, value in [('委托人', example['client_name']),
                          ('身份证号码', example['client_id']),
                          ('住址', example['client_address']),
                          ('联系电话', example['client_phone'])]:
        tabdef, vruns = _client_row(doc, prefix)
        assert len(vruns) == 1, f'{prefix}: 值 run 必须唯一，得到 {len(vruns)}'
        txt, ntabs, u, sz = _run_props(vruns[0])
        assert txt == value and ntabs == 2 and u == 'single', (prefix, txt, ntabs, u)
        assert sz == R.DEFAULT_SZ  # example 全短值 → 不缩字号
        assert tabdef[0][0] == 'center' and tabdef[-1] == ('right', 8400)

    # 律师费段：中文大写金额 + ¥ 后数字下划线；制式文字与 ¥ 本身无下划线
    fee_p = next(p for p in doc.paragraphs if p.text.startswith('一审阶段：'))
    seg = {(_run_props(r)[0], _run_props(r)[2]) for r in fee_p._p.findall(qn('w:r'))}
    assert ('叁万元', 'single') in seg
    assert ('30000', 'single') in seg
    assert all(u is None for t, u in seg if '¥' in t or t.startswith('一审阶段'))

    # contract_num 留空 → 渲染 5 个空格占位（律所系统回填）
    assert any('第M     号' in p.text for p in doc.paragraphs)


def test_render_long_address_shrinks():
    example = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    example['client_address'] = '北京市朝阳区示例路一号示例大厦示例单元示例楼层示例房间号示例园区'  # 30 CJK
    out, proc = _render(example, "long.docx")
    assert proc.returncode == 0, proc.stderr
    tabdef, vruns = _client_row(Document(str(out)), '住址')
    txt, ntabs, u, sz = _run_props(vruns[0])
    assert txt == example['client_address'] and ntabs == 2 and u == 'single'
    assert R.MIN_SZ <= sz < R.DEFAULT_SZ, f'长值必须缩字号，得到 sz={sz}'
    assert tabdef[0][0] == 'left'


def test_render_redline_warns_but_not_blocks():
    example = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    example['fee_clauses'] = [
        '一审阶段：在签订本合同之日，甲方向乙方支付前期固定律师服务费人民币贰万元（¥20000）。',
        '在实现委托事项目标后支付后期风险代理律师服务费，按实际回款金额的15%提成计算。',
    ]
    out, proc = _render(example, "redline.docx")
    assert proc.returncode == 0, 'WARN 不得拦截渲染'
    assert out.exists() and out.stat().st_size > 0
    assert 'WARN [fee-redline]' in proc.stderr
    assert 'fee_clauses[1]' in proc.stderr and '风险代理' in proc.stderr
    assert 'fee_clauses[0]' not in proc.stderr  # 干净段不误报


def test_render_example_no_warn():
    example = json.loads(EXAMPLE.read_text(encoding="utf-8"))
    out, proc = _render(example, "clean.docx")
    assert proc.returncode == 0 and out.exists()
    assert 'fee-redline' not in proc.stderr, proc.stderr


def test_golden_document_xml_byte_identical():
    """v1.4.2 引擎基线：卓建模板 + 固定日期 example → document.xml 逐字节一致。"""
    out, proc = _render(_fixed_example(), "golden.docx")
    assert proc.returncode == 0, proc.stderr
    got = _document_xml(out)
    expect = GOLDEN_XML.read_bytes()
    assert got == expect, f"document.xml 与 golden 不一致（len {len(got)} vs {len(expect)}）"


ROUNDTRIP_MAP = {
    "name": "民商事委托代理合同", "category": "litigation",
    "output_name": "民商事委托代理合同_{client_name}_V1.docx",
    "fields": [
        {"key": "COVER_DATE_CN", "label": "封面日期", "default": "$today_cn", "at": [{"para": 12, "text": "二〇二六年九月二日"}]},
        {"key": "YEAR", "label": "年份", "default": "$today_year", "at": [{"para": 15, "text": "2026"}]},
        {"key": "CLIENT_NAME", "label": "委托人", "required": True, "at": [{"para": 16, "text": "张三"}, {"para": 43, "text": "张三"}]},
        {"key": "CLIENT_ID", "label": "身份证", "required": True, "at": [{"para": 17, "text": "110101198001011234"}]},
        {"key": "CLIENT_ADDRESS", "label": "住址", "required": True, "at": [{"para": 18, "text": "北京市朝阳区示例路1号"}]},
        {"key": "CLIENT_PHONE", "label": "电话", "required": True, "at": [{"para": 19, "text": "13800138000"}]},
        {"key": "OPPONENT", "label": "对方", "required": True, "at": [{"para": 23, "text": "李四"}]},
        {"key": "CASE_MATTER", "label": "案由", "required": True, "at": [{"para": 23, "text": "民间借贷纠纷"}]},
        {"key": "LAWYER", "label": "律师", "default": "示例律师团队", "at": [{"para": 24, "text": "示例律师团队"}]},
        {"key": "PROC_CODES", "label": "程序", "required": True, "at": [{"para": 31, "text": "1、2"}]},
        {"key": "INVOICE_TYPE", "label": "发票", "default": "1", "at": [{"para": 41, "text": "1", "context": "类型为  1  "}]},
        {"key": "FEE_CLAUSES", "label": "律师费", "kind": "paragraphs", "underline": "auto", "at": [{"para": 34}, {"para": 35}]},
    ],
}


def test_init_roundtrip_on_zhuojian_layout():
    """已填合同 → init 建包 → 再渲染 → 与 golden 逐字节一致（初始化在真实版式上动态提炼变量/下划线/tab）。"""
    filled, proc = _render(_fixed_example(), "filled.docx")
    assert proc.returncode == 0, proc.stderr
    tdir = Path(tempfile.mkdtemp(prefix="retainer_rt_")) / "templates"
    r = _init(filled, ROUNDTRIP_MAP, tdir / "zj-rt")
    assert r.returncode == 0, r.stderr + r.stdout
    assert "CLIENT_NAME: 16 blank_row" in r.stdout and "FEE_CLAUSES 下划线模式：amounts" in r.stdout, r.stdout
    manifest = json.loads((tdir / "zj-rt" / "manifest.json").read_text(encoding="utf-8"))
    assert manifest["template_sha256"] and all("at" not in f for f in manifest["fields"])
    # manifest 不得保存任何样本值
    mtxt = json.dumps(manifest, ensure_ascii=False)
    for sample in ("张三", "110101198001011234", "13800138000", "李四"):
        assert sample not in mtxt, sample
    out, proc = _render_with(_fixed_example(), "rt.docx", tdir)
    assert proc.returncode == 0, proc.stderr
    assert _document_xml(out) == GOLDEN_XML.read_bytes(), "往返渲染与 golden 不一致"


def _build_generic_sample(path: Path):
    """非卓建版式的合成合同：tab leader 画线（值本身无下划线）/ 行内下划线无 tab / 金额下划线收费段 / 签章处重复姓名。"""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Twips
    d = Document()
    d.add_paragraph("委托代理合同")                                              # [0]
    p = d.add_paragraph()                                                        # [1] leader 行
    p.paragraph_format.tab_stops.add_tab_stop(Twips(7000), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)  # leader=underscore
    p.add_run("委托人：")
    p.add_run("\t王五")                                                         # 同一 run 内 [tab][王五]
    p2 = d.add_paragraph(); p2.add_run("联系电话：")                             # [2] 行内下划线
    r = p2.add_run("13900139000"); r.underline = True
    d.add_paragraph("承办律师：赵律师。")                                        # [3]
    d.add_paragraph("甲方因与钱七合同纠纷一案委托乙方。")                         # [4]
    p5 = d.add_paragraph(); p5.add_run("第一期：签约时支付人民币")               # [5] 收费段 1
    r = p5.add_run("壹万元"); r.underline = True
    p5.add_run("（¥"); r = p5.add_run("10000"); r.underline = True; p5.add_run("）。")
    p6 = d.add_paragraph(); p6.add_run("第二期：开庭前支付人民币")               # [6] 收费段 2
    r = p6.add_run("伍仟元"); r.underline = True
    p6.add_run("（¥"); r = p6.add_run("5000"); r.underline = True; p6.add_run("）。")
    d.add_paragraph("甲方（签章）：王五")                                        # [7] 残留
    d.save(str(path))


def test_init_generic_template_end_to_end():
    tmp = Path(tempfile.mkdtemp(prefix="retainer_generic_"))
    sample = tmp / "sample.docx"
    _build_generic_sample(sample)
    # inspect 能跑、能标出 leader 行
    ins = subprocess.run([sys.executable, str(INSPECT), str(sample)], capture_output=True, text=True)
    assert ins.returncode == 0 and "疑似空白下划线行" in ins.stdout and "leader=underscore" in ins.stdout, ins.stdout

    mp = {"name": "测试所委托合同", "category": "litigation", "fields": [
        {"key": "CLIENT_NAME", "label": "委托人", "required": True, "at": [{"para": 1, "text": "王五"}]},
        {"key": "CLIENT_PHONE", "label": "电话", "required": True, "at": [{"para": 2, "text": "13900139000"}]},
        {"key": "LAWYER", "label": "律师", "default": "赵律师", "at": [{"para": 3, "text": "赵律师"}]},
        {"key": "OPPONENT", "label": "对方", "required": True, "at": [{"para": 4, "text": "钱七"}]},
        {"key": "FEE_CLAUSES", "label": "律师费", "kind": "paragraphs", "at": [{"para": 5}, {"para": 6}]},
    ]}
    tdir = tmp / "templates"
    # 1) 签章处「王五」残留 → 隐私清扫拦截
    r = _init(sample, mp, tdir / "test-firm")
    assert r.returncode == 5 and "隐私清扫未通过" in r.stderr and "[7]" in r.stderr, (r.returncode, r.stderr)
    assert not (tdir / "test-firm" / "manifest.json").exists()
    # 2) 补 placement 后通过
    mp["fields"][0]["at"].append({"para": 7, "text": "王五"})
    r = _init(sample, mp, tdir / "test-firm")
    assert r.returncode == 0, r.stderr + r.stdout
    assert "CLIENT_NAME: 1 blank_row" in r.stdout and "CLIENT_PHONE: 2 inline" in r.stdout, r.stdout
    assert "FEE_CLAUSES 下划线模式：amounts" in r.stdout
    manifest = json.loads((tdir / "test-firm" / "manifest.json").read_text(encoding="utf-8"))
    assert manifest["slug"] == "test-firm" and manifest["fields"][-1]["underline"] == "amounts"
    skeleton = json.loads((tdir / "test-firm" / "contract.skeleton.json").read_text(encoding="utf-8"))
    assert skeleton["fee_clauses"] == [] and skeleton["client_name"] == ""
    # 3) 渲染
    data = {"client_name": "孙八", "client_phone": "13700137000", "lawyer": "", "opponent": "周九",
            "fee_clauses": ["第一期：签约时支付人民币贰万元（¥20000）。",
                            "第二期：开庭前支付人民币壹万元（¥10000）。",
                            "第三期：执行阶段另议。"]}
    out, proc = _render_with(data, "generic.docx", tdir)
    assert proc.returncode == 0, proc.stderr
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    # leader 行 → 单 run [tab][值][tab] 单下划线；tab stop 不再带 leader
    tabdef, vruns = _client_row(doc, "委托人")
    assert len(vruns) == 1
    txt, ntabs, u, sz = _run_props(vruns[0])
    assert (txt, ntabs, u) == ("孙八", 2, "single")
    assert tabdef[-1] == ("right", 7000) and tabdef[0][0] == "center"
    row_p = next(p for p in doc.paragraphs if p.text.startswith("委托人"))
    assert not any(t.get(qn("w:leader")) for t in row_p._p.find(qn("w:pPr")).find(qn("w:tabs")))
    # 行内下划线继承
    phone_p = next(p for p in doc.paragraphs if p.text.startswith("联系电话"))
    seg = {(_run_props(r)[0], _run_props(r)[2]) for r in phone_p._p.findall(qn("w:r"))}
    assert ("13700137000", "single") in seg
    # 默认值 / 对方 / 签章处
    assert "承办律师：赵律师。" in texts and "甲方因与周九合同纠纷一案委托乙方。" in texts and "甲方（签章）：孙八" in texts
    # 收费段 3 段、金额下划线
    fee = [p for p in doc.paragraphs if p.text.startswith(("第一期", "第二期", "第三期"))]
    assert [p.text for p in fee] == data["fee_clauses"]
    seg = {(_run_props(r)[0], _run_props(r)[2]) for r in fee[0]._p.findall(qn("w:r"))}
    assert ("贰万元", "single") in seg and ("20000", "single") in seg
    # 样本值零残留
    alltext = "\n".join(texts)
    for sample_val in ("王五", "13900139000", "钱七", "伍仟元", "5000"):
        assert sample_val not in alltext, sample_val


def test_no_pack_and_missing_required_errors():
    empty = Path(tempfile.mkdtemp(prefix="retainer_empty_"))
    out, proc = _render_with(_fixed_example(), "x.docx", empty)
    assert proc.returncode == 3 and "初始化" in proc.stderr and not out.exists(), (proc.returncode, proc.stderr)
    ex = _fixed_example(); ex["client_name"] = ""; ex["fee_clauses"] = []
    out, proc = _render(ex, "missing.docx")
    assert proc.returncode == 4 and "client_name" in proc.stderr and "fee_clauses" in proc.stderr, proc.stderr


def test_template_drift_warn_and_relock():
    import shutil
    tdir = Path(tempfile.mkdtemp(prefix="retainer_drift_")) / "templates"
    shutil.copytree(TEMPLATES_DIR / "zhuojian-minshang", tdir / "zhuojian-minshang")
    mpath = tdir / "zhuojian-minshang" / "manifest.json"
    m = json.loads(mpath.read_text(encoding="utf-8")); m["template_sha256"] = "0" * 64
    mpath.write_text(json.dumps(m, ensure_ascii=False), encoding="utf-8")
    out, proc = _render_with(_fixed_example(), "drift.docx", tdir)
    assert proc.returncode == 0 and "WARN [template-drift]" in proc.stderr and out.exists()
    r = subprocess.run([sys.executable, str(INIT), "relock", str(tdir / "zhuojian-minshang")], capture_output=True, text=True)
    assert r.returncode == 0, r.stderr
    out, proc = _render_with(_fixed_example(), "relocked.docx", tdir)
    assert proc.returncode == 0 and "template-drift" not in proc.stderr
    assert _document_xml(out) == GOLDEN_XML.read_bytes()


# 需要卓建模板包（私有，不随公开仓发布）的用例：公开仓里 templates/ 为空 → 跳过
HAS_ZJ_PACK = (TEMPLATES_DIR / "zhuojian-minshang" / "manifest.json").exists() and GOLDEN_XML.exists()

TESTS = [
    (test_num2cn_documented_examples, False),
    (test_redline_scan_positive, False),
    (test_redline_scan_negative, False),
    (test_render_example_layout, True),
    (test_render_long_address_shrinks, True),
    (test_render_redline_warns_but_not_blocks, True),
    (test_render_example_no_warn, True),
    (test_golden_document_xml_byte_identical, True),
    (test_init_roundtrip_on_zhuojian_layout, True),
    (test_init_generic_template_end_to_end, False),
    (test_no_pack_and_missing_required_errors, True),
    (test_template_drift_warn_and_relock, True),
]


def _pytest_guard():
    """pytest 收集时也按同一规则跳过。"""
    try:
        import pytest
    except ImportError:
        return
    for t, needs_pack in TESTS:
        if needs_pack and not HAS_ZJ_PACK:
            globals()[t.__name__] = pytest.mark.skip(reason="需要私有模板包 zhuojian-minshang")(t)


_pytest_guard()


def run():
    skipped = 0
    for t, needs_pack in TESTS:
        if needs_pack and not HAS_ZJ_PACK:
            print(f"[SKIP] {t.__name__}（需要私有模板包）")
            skipped += 1
            continue
        t()
        print(f"[OK] {t.__name__}")
    print(f"\nALL PASS（skipped {skipped}）" if skipped else "\nALL PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(run())
