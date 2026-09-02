#!/usr/bin/env python3
"""委托代理合同渲染引擎（v2.0，manifest 驱动）：contract.json + 模板包 → .docx

用法：
  render.py <contract.json> [<output.docx>] [--template <slug>] [--templates-dir <dir>]
  render.py --list-templates [--templates-dir <dir>]

模板包（template pack）= templates/<slug>/ 目录，含：
  template.docx   占位符版模板（{{KEY}} 形式），由 init_template.py 从使用者的合同样本生成
  manifest.json   字段清单 + 提问话术 + 默认值 + 渲染方式（同样由 init_template.py 生成）

manifest.json 结构：
{
  "manifest_version": 1,
  "slug": "lanhai-minshang",               # 目录名
  "name": "民商事委托代理合同",              # 展示名（多模板时 agent 用它问「用哪份合同」）
  "category": "litigation",                # litigation | criminal | advisory | other（决定 SKILL.md 哪些审查项适用）
  "template": "template.docx",
  "template_sha256": "…",                  # 模板指纹；不符 → stderr WARN（防漂移；`init_template.py relock` 重锁）
  "output_name": "民商事委托代理合同_{client_name}_V1.docx",   # 未指定输出路径时的默认文件名（可引用 json 字段）
  "fields": [
    {"key": "CLIENT_NAME",                 # 模板占位符 {{CLIENT_NAME}}
     "json": "client_name",                # contract.json 字段名（缺省 = key 小写）
     "label": "委托人姓名 / 单位全称",
     "ask": "委托人叫什么？个人给姓名，单位给全称。",   # agent 向使用者提问的话术（自然语言）
     "required": true,
     "default": "",                        # 空值回退；支持 $today_cn / $today_year 两个动态值
     "kind": "text"},                      # text（默认）| paragraphs（数组逐段成段，如收费条款）
    {"key": "FEE_CLAUSES", "json": "fee_clauses", "kind": "paragraphs",
     "underline": "amounts",               # amounts：中文大写金额与 ¥ 数字加下划线 | none：整段继承锚点段格式
     "redline_scan": true}                 # 渲染前跑收费条款红线词扫描（默认 true）
  ]
}

contract.json = { "<json 字段名>": 值, ..., "template": "<slug>"(可选), "output_path": "..."(可选) }

渲染规则（版式全部继承模板，引擎不写死任何律所的几何参数）：
- text 字段：{{KEY}} 所在 run 就地替换，继承该 run 的字体 / 字号 / 下划线（跨 run 的占位符合并到首 run）。
- 空白下划线行（结构判定，不靠标签文字）：{{KEY}} 独占一个带下划线的 run、且其前一个 run 是制表符 →
  整行重建为「label runs + 一个值 run（[tab][值][tab]，单条 run 下划线贯穿）」，短值 center-tab 居中、
  长值 left 布局自动缩字号；首行缩进 / 右侧 tab 位置 / 字号全部从该段自身的 pPr / rPr 读取。
  （v1.0.6 定型的单 run 单机制方案：根治 leader 与 run 下划线并存的双线 / 断线。）
- paragraphs 字段：找到整段恰为 {{KEY}} 的锚点段，按数组逐段克隆插入（继承锚点段 pPr / 首 run rPr）。
"""
from __future__ import annotations

import hashlib
import json
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

SKILL_NAME = 'retainer-agreement'
SKILL_VERSION = '2.0.1'

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = SCRIPT_DIR.parent / 'templates'

CN_DIGITS = '〇一二三四五六七八九'

# ── 收费条款红线词扫描（v1.4.0，Step 2.5-E 的代码级兜底）────────────────────
# 婚姻/继承案禁风险代理此前 100% 停在 prompt 层——Step 2.5 审查被跳过时渲染端毫无防线。
# 渲染前对每段 paragraphs 字段按下表 re.search 扫描，命中打 stderr WARN、**不拦截**
# （一般民商事案风险代理合法，终判权留律师）。直接增删本表即可调整口径；label 原样出现在 WARN 文案里。
FEE_REDLINE_PATTERNS = [
    ('风险代理',    r'风险代理'),
    ('减免承诺',    r'减免'),
    ('判不离不收',  r'判不离'),
    ('对方反悔不收', r'对方反悔'),
    ('按比例提成',  r'提成|按[^，。；]{0,8}比例'),
    ('结果挂钩收费', r'(离婚证|调解书|胜诉|离婚成功)[^，。；]{0,8}(才|再)[^，。；]{0,4}(收|付)'),
]

# ── 空白下划线行布局参数（v1.0.6 定型）───────────────────────────────────────
# 几何量（首行缩进 / 右 tab 位置 / 字号）v2.0 起不再是常量，逐段从模板自身读取；下面只剩布局策略常量。
LEFT_GAP = 240             # left 布局下 label 与值之间的下划线间隙 twip
CENTER_THRESHOLD = 0.85    # 值视觉宽度 ≤ 可用宽度×此值 → center；否则 left + 缩字号
LEFT_SAFETY = 0.93         # left 布局缩字号留余量，确保单行不溢出
DEFAULT_SZ = 28            # 占位 run 未声明字号时的回退（四号 14pt = 28 half-points）
MIN_SZ = 18                # 缩字号下限 9pt（再长则接受 Word 末位换行，属极端 case）

# 视觉宽度估算（环境无中文字体文件，按 em 估算）：
# CJK 一个全角 = 1em = 字号，sz28 下精确 280；ASCII 取偏高值，保证缩字号留足余量。
WIDTH_CJK = 280
WIDTH_ASCII = 155
WIDTH_SPACE = 80


# ═══════════════════════════════════════════════════════════════════════════
# 模板包
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class Pack:
    slug: str
    dir: Path
    manifest: dict

    @property
    def template_path(self) -> Path:
        return self.dir / self.manifest.get('template', 'template.docx')

    @property
    def name(self) -> str:
        return self.manifest.get('name') or self.slug


def sha256_of(path: Path) -> str:
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


def load_pack(pack_dir: Path) -> Pack:
    mpath = Path(pack_dir) / 'manifest.json'
    with open(mpath, encoding='utf-8') as f:
        manifest = json.load(f)
    fields = manifest.get('fields')
    if not isinstance(fields, list) or not fields:
        raise ValueError(f'{mpath}: manifest.fields 必须是非空数组')
    for fd in fields:
        if not fd.get('key'):
            raise ValueError(f'{mpath}: 每个 field 必须有 key')
        if fd.get('kind', 'text') not in ('text', 'paragraphs'):
            raise ValueError(f'{mpath}: field {fd["key"]} kind 只能是 text / paragraphs')
    slug = manifest.get('slug') or Path(pack_dir).name
    return Pack(slug=slug, dir=Path(pack_dir), manifest=manifest)


def list_packs(templates_dir: Path) -> list[Pack]:
    templates_dir = Path(templates_dir)
    if not templates_dir.is_dir():
        return []
    packs = []
    for d in sorted(templates_dir.iterdir()):
        if d.is_dir() and (d / 'manifest.json').exists():
            packs.append(load_pack(d))
    return packs


class PackResolutionError(RuntimeError):
    pass


def resolve_pack(templates_dir: Path, slug: str | None) -> Pack:
    """定位模板包：显式 slug > 唯一包自动选 > 报错（零包 / 多包未指定）。"""
    packs = list_packs(templates_dir)
    if slug:
        for p in packs:
            if p.slug == slug or p.name == slug:
                return p
        raise PackResolutionError(
            f'找不到模板包「{slug}」。可用：{", ".join(p.slug for p in packs) or "（无）"}')
    if not packs:
        raise PackResolutionError(
            f'{templates_dir} 下没有任何模板包（缺 <slug>/manifest.json）。'
            '本 skill 没有模板不能出件——请按 SKILL.md 第 0 步，让使用者提供一份常用的委托代理合同 Word，'
            '用 scripts/inspect_template.py + scripts/init_template.py 初始化第一个模板包。')
    if len(packs) > 1:
        raise PackResolutionError(
            '有多个模板包，请用 --template <slug> 或 contract.json 的 "template" 指定：'
            + ', '.join(f'{p.slug}（{p.name}）' for p in packs))
    return packs[0]


def warn_template_lock(pack: Pack) -> bool:
    """模板指纹校验：manifest.template_sha256 与实际不符 → stderr WARN（不拦截）。返回是否一致。"""
    expect = pack.manifest.get('template_sha256')
    if not expect:
        return True
    actual = sha256_of(pack.template_path)
    if actual != expect:
        print(f'⚠ WARN [template-drift] 模板包「{pack.slug}」的 template.docx 与 manifest 指纹不符'
              f'（期望 {expect[:10]}…，实际 {actual[:10]}…）。若是你自己在 Word 里改过模板，'
              f'运行 `init_template.py relock {pack.dir}` 重新锁定；否则请检查模板是否被换过。',
              file=sys.stderr)
        return False
    return True


# ═══════════════════════════════════════════════════════════════════════════
# 字段 / 默认值
# ═══════════════════════════════════════════════════════════════════════════

def date_to_cn(d: date) -> str:
    """Format date as '二〇二六年五月二十日'."""
    year_cn = ''.join(CN_DIGITS[int(c)] for c in str(d.year))
    return f'{year_cn}年{_num_to_cn_short(d.month)}月{_num_to_cn_short(d.day)}日'


def _num_to_cn_short(n: int) -> str:
    """Short Chinese numeral for month/day (1-31). 1→一, 10→十, 11→十一, 20→二十, 21→二十一."""
    if n < 10:
        return CN_DIGITS[n]
    if n == 10:
        return '十'
    if n < 20:
        return '十' + CN_DIGITS[n % 10]
    tens, ones = divmod(n, 10)
    return CN_DIGITS[tens] + '十' + (CN_DIGITS[ones] if ones else '')


def resolve_default(value, today: date | None = None):
    today = today or date.today()
    if value == '$today_cn':
        return date_to_cn(today)
    if value == '$today_year':
        return str(today.year)
    if value == '$today_iso':
        return today.isoformat()
    return value if value is not None else ''


def json_name(field: dict) -> str:
    return field.get('json') or field['key'].lower()


def build_mapping(text_fields: list[dict], data: dict, today: date | None = None):
    """text 字段 → {KEY: 字符串值}；空值取 default；返回 (mapping, 缺失的必填 json 名列表)。"""
    mapping, missing = {}, []
    for fd in text_fields:
        raw = data.get(json_name(fd), '')
        val = '' if raw is None else str(raw)
        if not val.strip():
            val = str(resolve_default(fd.get('default', ''), today))
        if fd.get('required') and not val.strip():
            missing.append(json_name(fd))
        mapping[fd['key']] = val
    return mapping, missing


# ═══════════════════════════════════════════════════════════════════════════
# 通用占位符替换（run 级，保格式）
# ═══════════════════════════════════════════════════════════════════════════

def replace_placeholders_in_runs(p, mapping):
    """Replace {{KEY}} in paragraph runs, preserving per-run formatting.

    Per-run swap when the placeholder is contained in a single run; for placeholders
    spanning runs, join the spanning runs and write the result into the first one.
    空白下划线行已由 render_blank_row 单独重建，此处不再涉及它们。
    """
    if not p.runs:
        return False
    tokens = {k: f'{{{{{k}}}}}' for k in mapping}
    runs = p.runs
    full_text = ''.join(r.text for r in runs)
    if not any(t in full_text for t in tokens.values()):
        return False

    # Fast path: per-run replacement
    for r in runs:
        rt = r.text
        for k, t in tokens.items():
            if t in rt:
                rt = rt.replace(t, str(mapping[k]))
        if rt != r.text:
            r.text = rt

    # Spanning case
    full_text = ''.join(r.text for r in runs)
    if any(t in full_text for t in tokens.values()):
        boundaries = []
        cursor = 0
        for r in runs:
            boundaries.append((cursor, cursor + len(r.text), r))
            cursor += len(r.text)
        for k, token in tokens.items():
            while token in full_text:
                idx = full_text.find(token)
                end = idx + len(token)
                spanned = [b for b in boundaries if b[0] < end and b[1] > idx]
                if not spanned:
                    break
                joined = ''.join(s[2].text for s in spanned)
                replaced = joined.replace(token, str(mapping[k]), 1)
                spanned[0][2].text = replaced
                for s in spanned[1:]:
                    s[2].text = ''
                full_text = ''.join(r.text for r in runs)
                boundaries = []
                cursor = 0
                for r in runs:
                    boundaries.append((cursor, cursor + len(r.text), r))
                    cursor += len(r.text)
    return True


def iter_all_paragraphs(doc):
    """正文 → 表格 → 页眉页脚，全部段落（顺序与 v1 一致，正文优先）。"""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer,
                     sec.even_page_header, sec.even_page_footer):
            # 只读已存在的部件：访问「链接到前一节」的页眉/页脚会让 python-docx 自动新建 part，
            # 污染 sectPr（多出 headerReference/footerReference）——v2.0 golden 对照抓出的坑
            if part.is_linked_to_previous:
                continue
            for p in part.paragraphs:
                yield p


# ═══════════════════════════════════════════════════════════════════════════
# 空白下划线行（v1.0.6 单 run 单机制；v2.0 几何量改从模板段落读取）
# ═══════════════════════════════════════════════════════════════════════════

def _run_text(r_el):
    return ''.join(t.text or '' for t in r_el.findall(qn('w:t')))


def _run_is_tab(r_el):
    return r_el.find(qn('w:tab')) is not None or _run_text(r_el) == '\t'


def _run_underlined(r_el):
    rPr = r_el.find(qn('w:rPr'))
    if rPr is None:
        return False
    u = rPr.find(qn('w:u'))
    return u is not None and u.get(qn('w:val')) not in (None, 'none')


def is_blank_row_paragraph(p_el, token) -> bool:
    """结构判定「空白下划线行」：token 独占一个带下划线的 run，且其前一个非空 run 是制表符。

    不看标签文字——任何律所的「委托人：____」「电话：____」只要是「标签 + tab + 下划线值」
    结构都命中；同一占位符出现在普通正文（如发票节「名称：{{CLIENT_NAME}}」）则不命中。
    """
    runs = p_el.findall(qn('w:r'))
    for i, r in enumerate(runs):
        if _run_text(r) != token:
            continue
        if not _run_underlined(r):
            return False
        j = i - 1
        while j >= 0:
            prev = runs[j]
            if _run_is_tab(prev):
                return True
            if _run_text(prev) == '':
                j -= 1
                continue
            return False
        return False
    return False


def _visual_width(text, sz=DEFAULT_SZ):
    """估算文本视觉宽度 twip（默认 sz28；其他字号按比例缩放）。"""
    scale = sz / DEFAULT_SZ
    w = 0
    for ch in text:
        if ord(ch) < 128:
            w += WIDTH_SPACE if ch == ' ' else WIDTH_ASCII
        else:
            w += WIDTH_CJK
    return w * scale


def _rpr_sz(rPr, fallback=DEFAULT_SZ):
    if rPr is not None:
        sz_el = rPr.find(qn('w:sz'))
        if sz_el is not None and sz_el.get(qn('w:val')):
            try:
                return int(sz_el.get(qn('w:val')))
            except ValueError:
                pass
    return fallback


def _run_visual_width(r_el):
    """估算一个 <w:r> 元素的视觉宽度（按其 rPr 里的 w:sz）。"""
    return _visual_width(_run_text(r_el), _rpr_sz(r_el.find(qn('w:rPr'))))


def _apply_sz(rPr, sz):
    """把 <w:sz>/<w:szCs> 设为 sz（half-points）；缺失则按 CT_RPr 顺序补建。"""
    val = str(sz)
    sz_el = rPr.find(qn('w:sz'))
    if sz_el is None:
        sz_el = OxmlElement('w:sz')
        u = rPr.find(qn('w:u'))
        if u is not None:
            u.addprevious(sz_el)
        else:
            rPr.append(sz_el)
    sz_el.set(qn('w:val'), val)
    szCs_el = rPr.find(qn('w:szCs'))
    if szCs_el is None:
        szCs_el = OxmlElement('w:szCs')
        sz_el.addnext(szCs_el)
    szCs_el.set(qn('w:val'), val)


def _apply_underline(rPr, underline):
    """开/关 run 级单下划线。值 run 传 True，下划线贯穿 [tab][文字][tab]。"""
    for u in list(rPr.findall(qn('w:u'))):
        rPr.remove(u)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        szCs = rPr.find(qn('w:szCs'))
        if szCs is not None:
            szCs.addnext(u)
        else:
            rPr.append(u)


def _make_client_value_run(base_rPr, value, sz):
    """造空白行的值 run：一个 <w:r> 内含 [tab][value][tab]，整体单下划线。

    核心：彻底放弃 tab leader——下划线只来自这一个 run 的 <w:u w:val="single">，
    连续贯穿前导制表符空白 + 值文字 + 尾随制表符空白，一种机制、无接缝、无双线。
    两个 <w:tab/> 按文档顺序绑定 pPr 的两个 tab stop（第一个 center/left、第二个 right）。
    """
    r = OxmlElement('w:r')
    rPr = deepcopy(base_rPr) if base_rPr is not None else OxmlElement('w:rPr')
    _apply_sz(rPr, sz)
    _apply_underline(rPr, True)
    r.append(rPr)
    r.append(OxmlElement('w:tab'))          # 前导制表符（run 下划线覆盖其空白）
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = value or ''
    r.append(t)
    r.append(OxmlElement('w:tab'))          # 尾随制表符（run 下划线续到右页边距）
    return r


def _set_row_tabs(pPr, tabs):
    """重设 pPr 的 <w:tabs>（tabs = [(val, pos), ...]）。tab stop 不带 leader——下划线由值 run 的 <w:u> 画。"""
    old = pPr.find(qn('w:tabs'))
    if old is not None:
        pPr.remove(old)
    tabs_el = OxmlElement('w:tabs')
    for val, pos in tabs:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), val)
        tab.set(qn('w:pos'), str(int(pos)))
        tabs_el.append(tab)
    pPr.insert(0, tabs_el)  # tabs 在 CT_PPr 中排在 adjustRightInd/spacing/ind 之前


def _row_geometry(pPr, text_width):
    """从段落自身读取几何量：行起点（left + firstLine 缩进）与右侧 tab 位置（最右 tab stop，缺省=版心右缘）。"""
    start = 0
    right_tab = int(text_width)
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            for attr in ('w:left', 'w:start', 'w:firstLine'):
                v = ind.get(qn(attr))
                if v:
                    try:
                        start += int(v)
                    except ValueError:
                        pass
        tabs = pPr.find(qn('w:tabs'))
        if tabs is not None:
            positions = []
            for t in tabs.findall(qn('w:tab')):
                try:
                    positions.append(int(t.get(qn('w:pos'))))
                except (TypeError, ValueError):
                    pass
            if positions:
                right_tab = max(positions)
    return start, right_tab


def _plan_client_row(value, label_end, right_tab, base_sz=DEFAULT_SZ):
    """决定空白行布局。返回 {'sz', 'tabs'}。

    - 值短（视觉宽度 ≤ 可用宽度×CENTER_THRESHOLD）→ center-tab 居中，原字号
    - 值长 → left 布局：值左对齐起于 label_end+LEFT_GAP，自动缩字号保证单行不溢出
    两种布局值 run 结构相同（一个 run 内 [tab][文字][tab]）；仅 pPr 的 tab 定义与字号不同。
    """
    value_w = _visual_width(value, base_sz)
    available = right_tab - label_end

    if value_w <= available * CENTER_THRESHOLD:
        midpoint = round((label_end + right_tab) / 2)
        return {'sz': base_sz,
                'tabs': [('center', midpoint), ('right', right_tab)]}

    value_start = label_end + LEFT_GAP
    avail_left = right_tab - value_start
    if value_w <= 0:
        sz = base_sz
    else:
        raw = base_sz * avail_left * LEFT_SAFETY / value_w
        sz = max(MIN_SZ, min(base_sz, int(raw)))
    return {'sz': sz,
            'tabs': [('left', value_start), ('right', right_tab)]}


def render_blank_row(p, token, value, text_width):
    """完全重建一行「标签 + 下划线空白」段落。

    保留 label run（不动其加粗/字号），删除占位符 run 及前后 tab run，按 _plan_client_row
    的布局重新写入 **一个** 值 run（内含 [tab][文字][tab]，整体单下划线）+ 对应 pPr tab 定义。
    返回 True 表示已重建；占位符跨 run（理论上不会）则返回 False 交给通用替换兜底。
    """
    p_el = p._p
    pPr = p_el.find(qn('w:pPr'))

    label_els = []
    ph_el = None
    for r_el in p_el.findall(qn('w:r')):
        txt = _run_text(r_el)
        if token in txt:
            ph_el = r_el
            break
        if _run_is_tab(r_el) and not txt.strip():
            continue  # 旧 tab run，丢弃
        label_els.append(r_el)
    if ph_el is None:
        return False

    base_rPr = ph_el.find(qn('w:rPr'))
    start, right_tab = _row_geometry(pPr, text_width)
    label_w = sum(_run_visual_width(r) for r in label_els)
    label_end = start + label_w
    plan = _plan_client_row(str(value), label_end, right_tab, _rpr_sz(base_rPr))

    # 清空 pPr 以外的全部子节点（旧 run / _GoBack 书签等）
    for child in list(p_el):
        if child is not pPr:
            p_el.remove(child)

    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    _set_row_tabs(pPr, plan['tabs'])

    for r_el in label_els:          # label run 原样放回
        p_el.append(r_el)
    # 值 [tab][文字][tab] 同属一个 run、共用一条 run 下划线（根治断线/双线）
    p_el.append(_make_client_value_run(base_rPr, str(value), plan['sz']))
    return True


# ═══════════════════════════════════════════════════════════════════════════
# 红线词扫描
# ═══════════════════════════════════════════════════════════════════════════

def scan_fee_redlines(fee_clauses):
    """对每段收费条款扫 FEE_REDLINE_PATTERNS。返回 [(条款下标, label, 命中文本), ...]。只扫描不拦截。"""
    hits = []
    for i, clause in enumerate(fee_clauses):
        for label, pat in FEE_REDLINE_PATTERNS:
            m = re.search(pat, str(clause))
            if m:
                hits.append((i, label, m.group()))
    return hits


def warn_fee_redlines(fee_clauses, field_name='fee_clauses'):
    """扫描并向 stderr 打 WARN；返回命中数。不 BLOCK——渲染继续，终判权留律师。"""
    hits = scan_fee_redlines(fee_clauses)
    for i, label, frag in hits:
        print(f'⚠ WARN [fee-redline] {field_name}[{i}] 命中「{label}」（…{frag}…）：'
              '婚姻/继承案属违法风险代理（须挂程序节点改写）；'
              '一般民商事案请复核上限/保底（Step 2.5-D）。', file=sys.stderr)
    return len(hits)


# ═══════════════════════════════════════════════════════════════════════════
# 多段字段（收费条款等）
# ═══════════════════════════════════════════════════════════════════════════

def find_anchor_paragraph(doc, token):
    """定位整段恰为 {{KEY}} 的锚点段。"""
    for p in doc.paragraphs:
        if p.text.strip() == token:
            return p
    raise RuntimeError(f'模板里找不到独占一段的 {token} 锚点')


# 金额下划线规则（underline=amounts）：
# - 大写中文金额（含可选「元」）：下划线
# - ¥ 后的阿拉伯数字（含逗号）：下划线
# - 其他制式文字 / 括号 / ¥ 符号本身：无下划线
CHINESE_AMOUNT_RE = re.compile(r'[壹贰叁肆伍陆柒捌玖拾佰仟万亿零]+元?')
ARABIC_AMOUNT_RE = re.compile(r'(?<=¥)[\d,]+')


def parse_fee_clause(text):
    """Return list of (text, underline) segments for a fee clause string."""
    spans = []
    for m in CHINESE_AMOUNT_RE.finditer(text):
        spans.append((m.start(), m.end()))
    for m in ARABIC_AMOUNT_RE.finditer(text):
        spans.append((m.start(), m.end()))
    spans.sort()
    merged = []
    for s, e in spans:
        if merged and s <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
        else:
            merged.append((s, e))
    segments = []
    pos = 0
    for s, e in merged:
        if s > pos:
            segments.append((text[pos:s], False))
        segments.append((text[s:e], True))
        pos = e
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments


def _clause_segments(text, underline_mode):
    if underline_mode == 'amounts':
        return parse_fee_clause(text)
    return [(text, None)]   # None = 继承锚点 run 的 rPr，不动其下划线


def _make_run_xml(template_rPr, text, underline):
    """Build a new <w:r> with template_rPr (deepcopy) and given text; underline True/False 覆盖、None 继承。"""
    r = etree.SubElement(etree.Element(qn('w:dummy')), qn('w:r'))
    if template_rPr is not None:
        rPr = deepcopy(template_rPr)
    elif underline:
        rPr = OxmlElement('w:rPr')   # 锚点 run 无 rPr（纯默认格式的模板）也要能画金额下划线
    else:
        rPr = None
    if rPr is not None:
        if underline is not None:
            for u in list(rPr.findall(qn('w:u'))):
                rPr.remove(u)
            if underline:
                u_el = etree.SubElement(rPr, qn('w:u'))
                u_el.set(qn('w:val'), 'single')
        r.append(rPr)
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return r


def render_clause_into_paragraph(paragraph, clause_text, underline_mode='amounts'):
    """Replace paragraph runs with multi-run rendering of a clause."""
    p_xml = paragraph._p
    first_r = p_xml.find(qn('w:r'))
    base_rPr = first_r.find(qn('w:rPr')) if first_r is not None else None
    for r in list(p_xml.findall(qn('w:r'))):
        p_xml.remove(r)
    for text, underline in _clause_segments(clause_text, underline_mode):
        if not text:
            continue
        p_xml.append(_make_run_xml(base_rPr, text, underline))


def insert_clause_paragraph_after(paragraph, clause_text, underline_mode='amounts'):
    """Insert a new paragraph after the given one, cloning paragraph format, render the clause."""
    new_p = deepcopy(paragraph._p)
    for r in list(new_p.findall(qn('w:r'))):
        new_p.remove(r)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    src_first_r = paragraph._p.find(qn('w:r'))
    base_rPr = src_first_r.find(qn('w:rPr')) if src_first_r is not None else None
    for text, underline in _clause_segments(clause_text, underline_mode):
        if not text:
            continue
        new_p.append(_make_run_xml(base_rPr, text, underline))
    return new_para


# 兼容旧名（v1.x 外部调用）
render_fee_clause_into_paragraph = render_clause_into_paragraph
insert_fee_paragraph_after = insert_clause_paragraph_after


# ═══════════════════════════════════════════════════════════════════════════
# 渲染主流程
# ═══════════════════════════════════════════════════════════════════════════

def _text_width(doc) -> int:
    sec = doc.sections[0]
    try:
        return int(sec.page_width.twips - sec.left_margin.twips - sec.right_margin.twips)
    except Exception:
        return 8400


def render_document(data: dict, pack: Pack, today: date | None = None):
    """渲染成 python-docx Document（不落盘）。"""
    manifest = pack.manifest
    fields = manifest['fields']
    text_fields = [f for f in fields if f.get('kind', 'text') == 'text']
    para_fields = [f for f in fields if f.get('kind') == 'paragraphs']

    mapping, missing = build_mapping(text_fields, data, today)
    for fd in para_fields:
        items = data.get(json_name(fd))
        if fd.get('required', True) and (not isinstance(items, list) or not items):
            missing.append(json_name(fd))
    if missing:
        raise ValueError('必填字段为空：' + ', '.join(missing) + '（回到 SKILL.md 第 1 步向使用者补问）')

    if not pack.template_path.exists():
        raise FileNotFoundError(f'Template missing: {pack.template_path}')
    doc = Document(str(pack.template_path))
    text_width = _text_width(doc)

    # 1. 空白下划线行：结构判定 → 整行重建
    for p in doc.paragraphs:
        txt = p.text
        for fd in text_fields:
            token = f'{{{{{fd["key"]}}}}}'
            if token in txt and is_blank_row_paragraph(p._p, token):
                render_blank_row(p, token, mapping[fd['key']], text_width)
                break

    # 2. 其余占位符：run 级就地替换（正文 / 表格 / 页眉页脚）
    for p in iter_all_paragraphs(doc):
        replace_placeholders_in_runs(p, mapping)

    # 3. 多段字段
    for fd in para_fields:
        items = data.get(json_name(fd)) or []
        if not items:
            continue
        items = [str(x) for x in items]
        if fd.get('redline_scan', True):
            warn_fee_redlines(items, json_name(fd))
        mode = fd.get('underline', 'amounts')
        anchor = find_anchor_paragraph(doc, f'{{{{{fd["key"]}}}}}')
        render_clause_into_paragraph(anchor, items[0], mode)
        for clause in items[1:]:
            anchor = insert_clause_paragraph_after(anchor, clause, mode)

    # 产物指纹（core.xml，正文零出现）：哪份模板包、哪个版本渲的
    try:
        doc.core_properties.category = (
            f'LDS {SKILL_NAME}/{SKILL_VERSION} tpl:{pack.slug}@{sha256_of(pack.template_path)[:10]}')
    except Exception:
        pass
    return doc


def default_output_path(data: dict, pack: Pack) -> Path:
    pattern = pack.manifest.get('output_name') or f'{pack.name}_{{client_name}}_V1.docx'
    class _Safe(dict):
        def __missing__(self, k):
            return ''
    name = pattern.format_map(_Safe({k: ('' if v is None else v) for k, v in data.items()}))
    return Path.cwd() / name


def render(contract_path: Path, output_path: Path | None = None,
           template_slug: str | None = None, templates_dir: Path | None = None) -> Path:
    with open(contract_path, encoding='utf-8') as f:
        data = json.load(f)
    pack = resolve_pack(templates_dir or TEMPLATES_DIR, template_slug or data.get('template'))
    warn_template_lock(pack)
    doc = render_document(data, pack)
    if output_path is None:
        op = data.get('output_path')
        output_path = Path(op) if op else default_output_path(data, pack)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'Rendered → {output_path}  [template: {pack.slug}]')
    return output_path


def _parse_args(argv):
    opts = {'template': None, 'templates_dir': None, 'list': False}
    positional = []
    it = iter(argv)
    for a in it:
        if a == '--template':
            opts['template'] = next(it, None)
        elif a.startswith('--template='):
            opts['template'] = a.split('=', 1)[1]
        elif a == '--templates-dir':
            opts['templates_dir'] = next(it, None)
        elif a.startswith('--templates-dir='):
            opts['templates_dir'] = a.split('=', 1)[1]
        elif a == '--list-templates':
            opts['list'] = True
        elif a in ('-h', '--help'):
            print(__doc__)
            sys.exit(0)
        else:
            positional.append(a)
    return opts, positional


def main():
    opts, positional = _parse_args(sys.argv[1:])
    templates_dir = Path(opts['templates_dir']) if opts['templates_dir'] else TEMPLATES_DIR
    if opts['list']:
        packs = list_packs(templates_dir)
        if not packs:
            print(f'（{templates_dir} 下没有模板包）')
        for p in packs:
            print(f'{p.slug}\t{p.name}\t{p.manifest.get("category", "")}')
        return
    if not positional:
        print('usage: render.py <contract.json> [<output.docx>] [--template <slug>]', file=sys.stderr)
        sys.exit(2)
    contract_path = Path(positional[0])
    output_path = Path(positional[1]) if len(positional) >= 2 else None
    try:
        render(contract_path, output_path, opts['template'], templates_dir)
    except PackResolutionError as e:
        print(f'✗ {e}', file=sys.stderr)
        sys.exit(3)
    except ValueError as e:
        print(f'✗ {e}', file=sys.stderr)
        sys.exit(4)


if __name__ == '__main__':
    main()
