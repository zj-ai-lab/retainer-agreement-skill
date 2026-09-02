#!/usr/bin/env python3
"""模板初始化：使用者的合同样本 .docx + agent 确认过的字段映射 map.json → 模板包 templates/<slug>/。

用法：
  init_template.py init --docx <合同样本.docx> --map <map.json> --out <templates/<slug>> [--allow-leftover] [--force]
  init_template.py relock <templates/<slug>>        # 在 Word 里改过模板后重新锁定指纹

map.json（agent 与使用者用自然语言确认后写出；使用者不需要看到它）：
{
  "name": "民商事委托代理合同",
  "category": "litigation",                 # litigation | criminal | advisory | other
  "output_name": "民商事委托代理合同_{client_name}_V1.docx",
  "fields": [
    {"key": "CLIENT_NAME", "label": "委托人姓名", "ask": "委托人叫什么？", "required": true,
     "at": [{"para": 16, "text": "张三"}, {"para": 42, "text": "张三"}]},
    {"key": "INVOICE_TYPE", "label": "发票类型", "default": "1",
     "at": [{"para": 40, "text": "1", "context": "类型为  1  （"}]},      # context：段内定位用的更长片段
    {"key": "FEE_CLAUSES", "label": "律师费条款", "kind": "paragraphs", "underline": "auto", "required": true,
     "at": [{"para": 34}, {"para": 35}]}                                   # 整段：首段成锚点，其余删除
  ]
}
定位方式：para = inspect_template.py 输出的正文段号；cell = [表, 行, 列, 段] 定位表格内段落。

做的事：
1. 把每处样本值替换成 {{KEY}}（run 级、保留原 run 的字体 / 字号 / 下划线；跨 run 时首 run 收前缀+占位、末 run 留后缀）
2. 动态识别「空白下划线行」（值带下划线 + 制表符，或 tab leader 画线）→ 规整成「label + tab run + 独占下划线 run」，
   渲染时按单 run 单机制方案重建整行（不写死任何标签文字 / 几何量，都从段落自身读）
3. paragraphs 字段：首段清空成锚点 {{KEY}}，其余段删除；underline=auto 时看样本里金额是否带下划线决定 amounts / none
4. 隐私清扫：所有样本值在整份文档（正文 / 表格 / 页眉页脚）里不得残留，残留即失败（--allow-leftover 放行）
5. 写 template.docx + manifest.json（含 template_sha256 指纹）+ contract.skeleton.json（contract.json 骨架）
样本原件不复制进模板包；manifest 不保存任何样本值。
"""
from __future__ import annotations

import json
import re
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from render import (ARABIC_AMOUNT_RE, CHINESE_AMOUNT_RE, _run_is_tab, _run_text,  # noqa: E402
                    _run_underlined, is_blank_row_paragraph, load_pack, sha256_of)

MANIFEST_VERSION = 1
LEADER_LINE = ('underscore', 'heavy', 'hyphen')


class InitError(RuntimeError):
    pass


# ── run 级文本编辑 ────────────────────────────────────────────────────────────

def _merge_ts(r_el):
    """把一个 run 里多个 <w:t> 合并成第一个（仅在需要编辑该 run 时调用）。"""
    ts = r_el.findall(qn('w:t'))
    if len(ts) <= 1:
        return ts[0] if ts else None
    ts[0].text = ''.join(t.text or '' for t in ts)
    ts[0].set(qn('xml:space'), 'preserve')
    for t in ts[1:]:
        r_el.remove(t)
    return ts[0]


def _set_run_text(r_el, text):
    t = _merge_ts(r_el)
    if t is None:
        t = OxmlElement('w:t')
        r_el.append(t)
    t.text = text
    t.set(qn('xml:space'), 'preserve')


def replace_in_paragraph(p_el, needle, replacement, context=None):
    """在段落 run 序列里把 needle 替换成 replacement（一次）。返回承载 replacement 的 run 元素。

    context 给出时先在段落全文里定位 context，再在其中找 needle（同一段落多处相同短值时用）。
    """
    runs = [r for r in p_el.iter(qn('w:r'))]
    texts = [_run_text(r) for r in runs]
    full = ''.join(texts)
    if context:
        cpos = full.find(context)
        if cpos < 0:
            raise InitError(f'段落里找不到 context「{context}」：{full[:60]!r}')
        rel = context.find(needle)
        if rel < 0:
            raise InitError(f'context「{context}」里不含 text「{needle}」')
        idx = cpos + rel
    else:
        idx = full.find(needle)
        if idx < 0:
            raise InitError(f'段落里找不到「{needle}」：{full[:80]!r}')
    end = idx + len(needle)

    bounds = []
    cursor = 0
    for r, t in zip(runs, texts):
        bounds.append((cursor, cursor + len(t), r))
        cursor += len(t)
    spanned = [(s, e, r) for s, e, r in bounds if s < end and e > idx and e > s]
    if not spanned:
        raise InitError(f'定位失败：「{needle}」')
    s0, e0, r0 = spanned[0]
    if len(spanned) == 1:
        t0 = texts[runs.index(r0)]
        _set_run_text(r0, t0[:idx - s0] + replacement + t0[end - s0:])
        return r0
    # 跨 run：首 run = 前缀 + replacement；中间 run 清空；末 run 留后缀
    t0 = texts[runs.index(r0)]
    _set_run_text(r0, t0[:idx - s0] + replacement)
    for s, e, r in spanned[1:-1]:
        _set_run_text(r, '')
    sl, el, rl = spanned[-1]
    tl = texts[runs.index(rl)]
    _set_run_text(rl, tl[end - sl:])
    return r0


def _split_text_with_tabs(text):
    """'a\\tb' → [('t','a'),('tab',None),('t','b')]，跳过空文本。"""
    parts = []
    for i, seg in enumerate(text.split('\t')):
        if i:
            parts.append(('tab', None))
        if seg:
            parts.append(('t', seg))
    return parts


def _new_run(rPr, kind, text=None):
    r = OxmlElement('w:r')
    if rPr is not None:
        r.append(deepcopy(rPr))
    if kind == 'tab':
        r.append(OxmlElement('w:tab'))
    else:
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
    return r


def isolate_token_run(r_el, token):
    """把承载 token 的 run 拆成：[前置子节点] [前缀文字] [token] [后缀文字] [后置子节点]，每段独立 run、同 rPr。

    目的：让 {{KEY}} 独占一个 run、其前的制表符成为独立 tab run —— render 的结构判定正好命中。
    返回 token run。
    """
    rPr = r_el.find(qn('w:rPr'))
    children = [c for c in r_el if c is not rPr]
    t_els = [c for c in children if c.tag == qn('w:t')]
    t_el = t_els[0] if t_els else None
    if t_el is None or token not in (t_el.text or ''):
        return r_el
    txt = t_el.text or ''
    k = txt.find(token)
    prefix, suffix = txt[:k], txt[k + len(token):]
    pre_children = children[:children.index(t_el)]
    post_children = children[children.index(t_el) + 1:]
    if not pre_children and not post_children and not prefix and not suffix:
        return r_el  # 已独占

    new_runs = []
    for c in pre_children:
        nr = OxmlElement('w:r')
        if rPr is not None:
            nr.append(deepcopy(rPr))
        nr.append(deepcopy(c))
        new_runs.append(nr)
    for kind, text in _split_text_with_tabs(prefix):
        new_runs.append(_new_run(rPr, kind, text))
    token_run = _new_run(rPr, 't', token)
    new_runs.append(token_run)
    for kind, text in _split_text_with_tabs(suffix):
        new_runs.append(_new_run(rPr, kind, text))
    for c in post_children:
        nr = OxmlElement('w:r')
        if rPr is not None:
            nr.append(deepcopy(rPr))
        nr.append(deepcopy(c))
        new_runs.append(nr)
    parent = r_el.getparent()
    anchor = r_el
    for nr in new_runs:
        anchor.addnext(nr)
        anchor = nr
    parent.remove(r_el)
    return token_run


def _ensure_underline(r_el):
    rPr = r_el.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r_el.insert(0, rPr)
    for u in list(rPr.findall(qn('w:u'))):
        rPr.remove(u)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    szCs = rPr.find(qn('w:szCs'))
    if szCs is not None:
        szCs.addnext(u)
    else:
        rPr.append(u)


def _paragraph_has_leader(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        return False
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        return False
    return any(t.get(qn('w:leader')) in LEADER_LINE for t in tabs.findall(qn('w:tab')))


def _paragraph_has_tab(p_el):
    return any(_run_is_tab(r) for r in p_el.iter(qn('w:r')))


def normalize_blank_row(p_el, r_el, token):
    """动态识别空白下划线行并规整成 render 可结构判定的形态。返回 'blank_row' / 'inline'。

    判定：承载值的 run 带下划线 且 段内有制表符；或段落 tab stop 带 leader（用 leader 画线的模板）且值前有制表符。
    规整：token 独占 run + 前置 tab 独立 run；leader 画线的补上 run 下划线（渲染时 leader 会被去掉、由 run 下划线接管）。
    """
    underlined = _run_underlined(r_el)
    has_tab = _paragraph_has_tab(p_el) or '\t' in ''.join(_run_text(r) for r in p_el.iter(qn('w:r')))
    leader = _paragraph_has_leader(p_el)
    if not has_tab or not (underlined or leader):
        return 'inline'
    token_run = isolate_token_run(r_el, token)
    if not _run_underlined(token_run):
        _ensure_underline(token_run)
    return 'blank_row' if is_blank_row_paragraph(p_el, token) else 'inline'


# ── paragraphs 字段 ───────────────────────────────────────────────────────────

def _amounts_underlined(p_els):
    for p_el in p_els:
        for r in p_el.iter(qn('w:r')):
            if not _run_underlined(r):
                continue
            txt = _run_text(r)
            if CHINESE_AMOUNT_RE.search(txt) or re.search(r'\d', txt):
                return True
    return False


def make_anchor(p_el, token):
    """整段清成一个 {{KEY}} run（rPr 取原首个非空 run；下划线剥掉——渲染时按 underline 模式重画）。"""
    first_rPr = None
    for r in p_el.iter(qn('w:r')):
        if _run_text(r).strip():
            first_rPr = r.find(qn('w:rPr'))
            break
    if first_rPr is None:
        r0 = p_el.find(qn('w:r'))
        first_rPr = r0.find(qn('w:rPr')) if r0 is not None else None
    rPr = deepcopy(first_rPr) if first_rPr is not None else None
    if rPr is not None:
        for u in list(rPr.findall(qn('w:u'))):
            rPr.remove(u)
    pPr = p_el.find(qn('w:pPr'))
    for child in list(p_el):
        if child is not pPr:
            p_el.remove(child)
    p_el.append(_new_run(rPr, 't', token))


# ── 定位 / 清扫 ───────────────────────────────────────────────────────────────

def locate(doc, at):
    if 'para' in at:
        try:
            return doc.paragraphs[int(at['para'])]._p
        except IndexError:
            raise InitError(f'正文没有第 {at["para"]} 段（共 {len(doc.paragraphs)} 段）')
    if 'cell' in at:
        t, r, c, p = at['cell']
        try:
            return doc.tables[t].rows[r].cells[c].paragraphs[p]._p
        except IndexError:
            raise InitError(f'表格定位失败：{at["cell"]}')
    raise InitError(f'placement 缺 para / cell：{at}')


def all_text_units(doc):
    for i, p in enumerate(doc.paragraphs):
        yield f'[{i}]', p.text
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield f'[T{ti}.R{ri}.C{ci}.P{pi}]', p.text
    for si, sec in enumerate(doc.sections):
        for label, part in (('页眉', sec.header), ('页脚', sec.footer)):
            if part.is_linked_to_previous:
                continue
            for p in part.paragraphs:
                yield f'[{label}{si}]', p.text


ID_RE = re.compile(r'(?<!\d)\d{17}[\dXx](?!\d)')
MOBILE_RE = re.compile(r'(?<!\d)1[3-9]\d{9}(?!\d)')


def sweep(doc, samples):
    leftovers, suspicious = [], []
    for loc, text in all_text_units(doc):
        for s in samples:
            if len(s) >= 2 and s in text:
                leftovers.append((loc, s, text[:60]))
        for pat, what in ((ID_RE, '18 位身份证号'), (MOBILE_RE, '11 位手机号')):
            if pat.search(text):
                suspicious.append((loc, what, text[:60]))
    return leftovers, suspicious


# ── 主流程 ────────────────────────────────────────────────────────────────────

def init_pack(docx_path: Path, map_path: Path, out_dir: Path, allow_leftover=False, force=False):
    with open(map_path, encoding='utf-8') as f:
        mp = json.load(f)
    fields = mp.get('fields') or []
    if not fields:
        raise InitError('map.fields 为空')
    out_dir = Path(out_dir)
    if out_dir.exists() and any(out_dir.iterdir()) and not force:
        raise InitError(f'{out_dir} 已存在且非空；确认覆盖请加 --force')

    doc = Document(str(docx_path))
    report = {'placements': [], 'blank_rows': [], 'deleted_paragraphs': 0, 'underline_modes': {}}
    samples = []
    # 先把所有定位解析成元素（后面会删段，索引会漂）
    resolved = []
    for fd in fields:
        key = fd['key']
        if not re.fullmatch(r'[A-Z][A-Z0-9_]*', key):
            raise InitError(f'key 只能用大写字母数字下划线：{key}')
        ats = fd.get('at') or []
        if not ats:
            raise InitError(f'field {key} 没有 at 定位')
        resolved.append((fd, [(at, locate(doc, at)) for at in ats]))

    text_fields, para_fields, to_delete = [], [], []
    for fd, places in resolved:
        token = f'{{{{{fd["key"]}}}}}'
        if fd.get('kind', 'text') == 'paragraphs':
            para_fields.append(fd)
            p_els = [p for _, p in places]
            mode = fd.get('underline', 'auto')
            if mode == 'auto':
                mode = 'amounts' if _amounts_underlined(p_els) else 'none'
            fd['underline'] = mode
            report['underline_modes'][fd['key']] = mode
            make_anchor(p_els[0], token)
            to_delete.extend(p_els[1:])          # 删段推迟到隐私清扫之后：清扫报告的段号要与 inspect 编号一致
            report['placements'].append((fd['key'], f'{len(p_els)} 段 → 锚点 + 逐段插入'))
            continue
        text_fields.append(fd)
        for at, p_el in places:
            needle = at.get('text')
            if not needle:
                raise InitError(f'field {fd["key"]} 的 placement 缺 text：{at}')
            samples.append(needle)
            r_el = replace_in_paragraph(p_el, needle, token, at.get('context'))
            layout = normalize_blank_row(p_el, r_el, token)
            loc = at.get('para', at.get('cell'))
            report['placements'].append((fd['key'], f'{loc} {layout}'))
            if layout == 'blank_row':
                report['blank_rows'].append((fd['key'], loc))

    leftovers, suspicious = sweep(doc, sorted(set(samples), key=len, reverse=True))
    report['leftovers'] = leftovers
    report['suspicious'] = suspicious
    if leftovers and not allow_leftover:
        msg = '\n'.join(f'  {loc} 残留「{s}」：{t}' for loc, s, t in leftovers)
        raise InitError('隐私清扫未通过——样本值仍出现在这些位置（要么加 placement 把它也变成变量，'
                        '要么确认它是固定文字后加 --allow-leftover）：\n' + msg)

    for p_el in to_delete:
        p_el.getparent().remove(p_el)
        report['deleted_paragraphs'] += 1

    out_dir.mkdir(parents=True, exist_ok=True)
    tpl = out_dir / 'template.docx'
    doc.save(str(tpl))

    def _clean(fd):
        d = {k: v for k, v in fd.items() if k != 'at'}
        d.setdefault('json', fd['key'].lower())
        d.setdefault('kind', 'text')
        if d['kind'] == 'paragraphs':
            d.setdefault('redline_scan', True)
            d.setdefault('required', True)
        return d

    manifest = {
        'manifest_version': MANIFEST_VERSION,
        'slug': mp.get('slug') or out_dir.name,
        'name': mp.get('name') or out_dir.name,
        'category': mp.get('category', 'other'),
        'template': 'template.docx',
        'template_sha256': sha256_of(tpl),
        'output_name': mp.get('output_name') or f'{mp.get("name") or out_dir.name}_{{{text_fields[0]["key"].lower() if text_fields else "client_name"}}}_V1.docx',
        'created': date.today().isoformat(),
        'notes': mp.get('notes', ''),
        'fields': [_clean(fd) for fd in fields],
    }
    with open(out_dir / 'manifest.json', 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
        f.write('\n')

    skeleton = {'template': manifest['slug']}
    for fd in manifest['fields']:
        skeleton[fd['json']] = [] if fd['kind'] == 'paragraphs' else ''
    skeleton['output_path'] = ''
    skeleton['_labels'] = {fd['json']: fd.get('label', '') for fd in manifest['fields']}
    with open(out_dir / 'contract.skeleton.json', 'w', encoding='utf-8') as f:
        json.dump(skeleton, f, ensure_ascii=False, indent=2)
        f.write('\n')
    return manifest, report


def print_report(manifest, report, out_dir):
    print(f'✔ 模板包已生成：{out_dir}')
    print(f'  slug={manifest["slug"]}  name={manifest["name"]}  category={manifest["category"]}  fields={len(manifest["fields"])}')
    for key, what in report['placements']:
        print(f'  · {key}: {what}')
    if report['blank_rows']:
        print('  空白下划线行（渲染时整行重建、单 run 下划线）：' + ', '.join(f'{k}@{loc}' for k, loc in report['blank_rows']))
    for k, m in report['underline_modes'].items():
        print(f'  {k} 下划线模式：{m}')
    if report['deleted_paragraphs']:
        print(f'  删除样本多余段落 {report["deleted_paragraphs"]} 段')
    if report['leftovers']:
        print('  ⚠ 放行的残留样本值：' + '; '.join(f'{loc}「{s}」' for loc, s, _ in report['leftovers']))
    if report['suspicious']:
        print('  ⚠ 模板里仍有疑似证件号 / 手机号（请确认是律所固定信息而非当事人）：'
              + '; '.join(f'{loc} {w}' for loc, w, _ in report['suspicious']))
    print('  下一步：用 contract.skeleton.json 填一份示例数据（张三 / 李四），render.py 干跑一份让使用者在 Word 里看。')


def relock(pack_dir: Path):
    pack = load_pack(pack_dir)
    new = sha256_of(pack.template_path)
    old = pack.manifest.get('template_sha256')
    pack.manifest['template_sha256'] = new
    with open(Path(pack_dir) / 'manifest.json', 'w', encoding='utf-8') as f:
        json.dump(pack.manifest, f, ensure_ascii=False, indent=2)
        f.write('\n')
    print(f'✔ relock {pack.slug}: {(old or "")[:10]}… → {new[:10]}…')


def main():
    argv = sys.argv[1:]
    if not argv or argv[0] in ('-h', '--help'):
        print(__doc__)
        sys.exit(2)
    cmd, rest = argv[0], argv[1:]
    try:
        if cmd == 'relock':
            if not rest:
                raise InitError('relock 需要模板包目录')
            relock(Path(rest[0]))
            return
        if cmd != 'init':
            raise InitError(f'未知子命令 {cmd}')
        opts = {'docx': None, 'map': None, 'out': None}
        allow_leftover = force = False
        it = iter(rest)
        for a in it:
            if a in ('--docx', '--map', '--out'):
                opts[a[2:]] = next(it, None)
            elif a == '--allow-leftover':
                allow_leftover = True
            elif a == '--force':
                force = True
            else:
                raise InitError(f'未知参数 {a}')
        if not all(opts.values()):
            raise InitError('init 需要 --docx --map --out 三个参数')
        manifest, report = init_pack(Path(opts['docx']), Path(opts['map']), Path(opts['out']),
                                     allow_leftover=allow_leftover, force=force)
        print_report(manifest, report, opts['out'])
    except InitError as e:
        print(f'✗ {e}', file=sys.stderr)
        sys.exit(5)


if __name__ == '__main__':
    main()
