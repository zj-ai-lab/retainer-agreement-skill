#!/usr/bin/env python3
"""模板检视：把使用者给的合同 Word 逐段编号吐出来，附下划线 / 制表符 / 已有占位符标注，供 agent 判断哪些地方是变量。

用法：
  inspect_template.py <合同.docx> [--json]

文本输出每行：
  [16] 委托人（以下简称“甲方”）：⇥张三⇥  ‖ 下划线「张三」 ‖ tab停靠 center@6200 right@8400 ‖ 首行缩进562 ‖ 疑似空白下划线行
标注含义：
  ⇥                 制表符
  下划线「…」        该段里带下划线的文字（多个用 / 隔开）；「⇥」表示下划线覆盖制表符空白
  tab停靠            段落自带的 tab stop（val@pos，带 leader 时标 leader）
  疑似空白下划线行    「标签 + 制表符 + 下划线值」结构（或 tab leader 画线）——init 会按单 run 下划线方案重建
  占位符             已含 {{KEY}} 形式的占位符（说明这份 Word 已经是模板）
表格段落编号 [T0.R1.C2.P0]；页眉页脚 [页眉] / [页脚]（只列文本，init 目前不在页眉页脚里放变量）。
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from render import _run_is_tab, _run_text, _run_underlined  # noqa: E402

PLACEHOLDER_RE = re.compile(r'\{\{[A-Z0-9_]+\}\}')


def _tabstops(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        return [], None
    out = []
    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None:
        for t in tabs.findall(qn('w:tab')):
            out.append({'val': t.get(qn('w:val')), 'pos': t.get(qn('w:pos')), 'leader': t.get(qn('w:leader'))})
    ind = pPr.find(qn('w:ind'))
    first_line = ind.get(qn('w:firstLine')) if ind is not None else None
    return out, first_line


def describe_paragraph(p_el):
    """返回该段落的结构描述 dict。"""
    runs = p_el.findall(qn('w:r'))
    display = []
    underlined = []
    has_tab = False
    for r in runs:
        txt = _run_text(r)
        u = _run_underlined(r)
        piece = ''
        for child in r:
            tag = child.tag.split('}')[1]
            if tag == 't':
                piece += (child.text or '').replace('\t', '⇥')
            elif tag == 'tab':
                piece += '⇥'
        if '⇥' in piece:
            has_tab = True
        display.append(piece)
        if u and piece.strip('⇥ ') != '':
            underlined.append(piece)
        elif u and piece:
            underlined.append(piece)  # 只有 tab 的下划线 run（空白线）
    text = ''.join(display)
    tabstops, first_line = _tabstops(p_el)
    leader_line = any(t.get('leader') in ('underscore', 'heavy', 'hyphen') for t in tabstops)
    # 疑似空白下划线行：有制表符，且（值带下划线 或 tab leader 画线）
    blank_row = has_tab and (bool(underlined) or leader_line)
    return {
        'text': text,
        'underlined': underlined,
        'has_tab': has_tab,
        'tabstops': tabstops,
        'first_line': first_line,
        'blank_row': blank_row,
        'placeholders': PLACEHOLDER_RE.findall(text),
    }


def inspect(docx_path: Path):
    doc = Document(str(docx_path))
    entries = []
    for i, p in enumerate(doc.paragraphs):
        d = describe_paragraph(p._p)
        d['loc'] = f'[{i}]'
        d['para'] = i
        entries.append(d)
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    d = describe_paragraph(p._p)
                    d['loc'] = f'[T{ti}.R{ri}.C{ci}.P{pi}]'
                    d['cell'] = [ti, ri, ci, pi]
                    entries.append(d)
    for si, sec in enumerate(doc.sections):
        for label, part in (('页眉', sec.header), ('页脚', sec.footer)):
            if part.is_linked_to_previous:
                continue
            for p in part.paragraphs:
                if p.text.strip():
                    d = describe_paragraph(p._p)
                    d['loc'] = f'[{label}{si}]'
                    entries.append(d)
    return entries


def format_text(entries):
    lines = []
    for d in entries:
        if not d['text'].strip('⇥ ') and not d['underlined']:
            continue
        notes = []
        if d['underlined']:
            notes.append('下划线' + ''.join(f'「{u}」' for u in d['underlined']))
        if d['tabstops']:
            notes.append('tab停靠 ' + ' '.join(
                f"{t['val']}@{t['pos']}" + (f"(leader={t['leader']})" if t.get('leader') else '')
                for t in d['tabstops']))
        if d['first_line']:
            notes.append(f"首行缩进{d['first_line']}")
        if d['blank_row']:
            notes.append('疑似空白下划线行')
        if d['placeholders']:
            notes.append('占位符 ' + ' '.join(d['placeholders']))
        line = f"{d['loc']} {d['text']}"
        if notes:
            line += '  ‖ ' + ' ‖ '.join(notes)
        lines.append(line)
    return '\n'.join(lines)


def main():
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    as_json = '--json' in sys.argv
    if not args:
        print(__doc__)
        sys.exit(2)
    entries = inspect(Path(args[0]))
    if as_json:
        print(json.dumps(entries, ensure_ascii=False, indent=1))
    else:
        print(format_text(entries))
        n_ph = sum(len(d['placeholders']) for d in entries)
        n_blank = sum(1 for d in entries if d['blank_row'])
        print(f'\n—— 共 {len(entries)} 段；疑似空白下划线行 {n_blank} 段；已有占位符 {n_ph} 处'
              + ('（这份 Word 已经是模板，可直接按占位符建 manifest）' if n_ph else ''))


if __name__ == '__main__':
    main()
