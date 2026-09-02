#!/usr/bin/env python3
"""生成一份**虚构**的委托代理合同 Word（示例律师事务所 / 张三 / 李四），给没有现成模板的人试跑第 0 步初始化用。

用法：
  python3 examples/make_demo_sample.py [输出路径.docx]     # 默认 ./示例合同样本（虚构）.docx

样本刻意包含三种常见版式：当事人信息「标签 + 制表符 + 下划线值」空白行、行内下划线（承办律师）、
律师费分期段落且金额带下划线。所有内容均为虚构，与任何真实律所、当事人无关。
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt, Twips


def _cn_font(run, size=14, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '仿宋')


def _blank_row(doc, label, value, center_pos):
    """当事人信息行：label + [tab] + 下划线值 + [tab]，tab stop 带 underscore leader。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Twips(560)
    p.paragraph_format.tab_stops.add_tab_stop(Twips(center_pos), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.LINES)
    p.paragraph_format.tab_stops.add_tab_stop(Twips(8400), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)
    _cn_font(p.add_run(label))
    r = p.add_run('\t' + value + '\t')
    _cn_font(r)
    r.underline = True
    return p


def _para(doc, text, indent=True, size=14, bold=False, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Twips(560)
    if align:
        p.alignment = align
    _cn_font(p.add_run(text), size, bold)
    return p


def _fee(doc, prefix, cn_amount, digits, suffix='）。'):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Twips(560)
    _cn_font(p.add_run(prefix))
    r = p.add_run(cn_amount); _cn_font(r); r.underline = True
    _cn_font(p.add_run('（¥'))
    r = p.add_run(digits); _cn_font(r); r.underline = True
    _cn_font(p.add_run(suffix))
    return p


def build(path: Path):
    doc = Document()
    _para(doc, '委托代理合同', indent=False, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, '合同编号：示例字第 2026-001 号', indent=False, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _blank_row(doc, '委托人（甲方）：', '张三', 5200)
    _blank_row(doc, '身份证号码：', '110101199001011234', 4800)
    _blank_row(doc, '住址：', '北京市朝阳区示例路 1 号', 4200)
    _blank_row(doc, '联系电话：', '13800138000', 4600)
    _para(doc, '受托人（乙方）：示例律师事务所')
    _para(doc, '甲方因与李四民间借贷纠纷一案，委托乙方代理，双方经协商达成如下条款：')
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Twips(560)
    _cn_font(p.add_run('第一条  乙方指派本所'))
    r = p.add_run('王律师'); _cn_font(r); r.underline = True
    _cn_font(p.add_run('担任甲方的委托代理人，代理权限以授权委托书为准。'))
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Twips(560)
    _cn_font(p.add_run('第二条  甲方委托乙方代理的法律程序为：'))
    r = p.add_run(' 一审 '); _cn_font(r); r.underline = True
    _cn_font(p.add_run('（一审 / 二审 / 执行）。'))
    _para(doc, '第三条  甲方按如下方式向乙方支付律师服务费：')
    _fee(doc, '一审阶段：在签订本合同之日，甲方向乙方支付律师代理费人民币', '贰万元', '20000')
    _fee(doc, '开庭前三日内，甲方向乙方支付律师代理费人民币', '壹万元', '10000')
    _para(doc, '第四条  甲方应如实陈述案情并提供证据材料；因甲方隐瞒或虚构事实造成的不利后果由甲方承担。')
    _para(doc, '第五条  乙方不对案件结果作任何承诺或保证。')
    _para(doc, '第六条  诉讼费、保全费、鉴定费等第三方费用由甲方另行承担。')
    _para(doc, '第七条  本合同自双方签字盖章之日起生效，一式两份，双方各执一份。')
    _para(doc, '')
    _para(doc, '甲方（签章）：张三', indent=False)
    _para(doc, '乙方（签章）：示例律师事务所', indent=False)
    _para(doc, '二〇二六年  月  日', indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.save(str(path))
    return path


def main():
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('示例合同样本（虚构）.docx')
    build(out)
    print(f'已生成虚构示例合同：{out}')


if __name__ == '__main__':
    main()
